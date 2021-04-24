import pymysql.cursors
import docx
import re
from  pathlib import Path



def preproccess_file(file):
    '''文件预处理'''
    #对文件内容预处理(把文本集中到一个字符串中）
    paragraph_sum = len(file.paragraphs)

    para_list = []
    for idx, para in enumerate(file.paragraphs) :
        para_list.append(para)
    return para_list

def extract(keyword, para_list):
    '''按关键词提取内容'''
    #使用正则提取关键字后面的数字
    is_reference = False
    refer_list = []
    for idx, para in enumerate(para_list):
        if not is_reference:
            result = re.match('{}'.format(keyword), para.text)
            if result:
                is_reference = True
        else:
            refer_list.append(para);

    return refer_list

def check(refer_list):
    print("对参考文件进行格式检查")
    # 对引用部分提取匹配
    info_list = []
    # 提取基本格式，并检查
    for para in refer_list:
        # print statistical data;
        match = re.match('^\[(\d+)\].*', para.text)
        # print(para.text)
        if match:
            res = re.findall('\[(\w)\]', para.text) # 提取[]中数字和文本类别
            info_list.append(res)
        else:
            print("引用开头格式有误，内容为：", para.text)


    # 次序检查，期刊类别检查
    old = 0
    for idx, info in enumerate(info_list):
        if int(info[0]) != old + 1:
            print("引用顺序有误:", refer_list[idx].text)
        old = int(info[0])
        if len(info) > 1:
            if(info[1] == 'J'):
                print("引用了期刊：", refer_list[idx].text)
        #TODO 其他检查

    # print(result)
    # return result



def main():
    path = Path('./test.docx')
    keyword = '参考文献'
    file = docx.Document(path)
    file_text = preproccess_file(file) #
    reference = extract(keyword, file_text)
    check(reference)

if __name__ == '__main__':
    main()